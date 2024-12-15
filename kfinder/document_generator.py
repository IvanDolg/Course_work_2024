from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import logging
from typing import List

from kcatalog.models import Belmarc

logger = logging.getLogger(__name__)



def generate_bulletin(books: List[Belmarc], library: str, chapter_name: str):
    """
    Генерирует бюллетень новых поступлений в формате docx
    
    Args:
        books: список записей в формате Belmarc
        year: год бюллетеня
        month: месяц бюллетеня
    
    Returns:
        document: объект Document из python-docx
    """
    try:
        # Создаем новый документ
        doc = Document()
        
        # Устанавливаем стиль для всего документа
        style = doc.styles['Normal']
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        
        # Настраиваем параметры страницы
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Добавляем заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(f'Новые поступления литературы В {library}\n Тема: {chapter_name}')
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Добавляем пустую строку после заголовка
        doc.add_paragraph()
        
        # Добавляем информацию о каждой книге
        for book in books:
            # Создаем новый параграф с отступом первой строки
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(0.5)
            
            # Добавляем стиль для всего параграфа
            paragraph_style = paragraph._element.get_or_add_pPr()
            spacing = paragraph_style.get_or_add_spacing()
            spacing.set('line', '240')  # 240 twips = exactly 12pt
            spacing.set('lineRule', 'exact')
            
            # Безопасное получение значений с проверкой на None и обработка списков
            author_name = book.personal_name_subject.entry_element if book.personal_name_subject else ""
            author_run = paragraph.add_run(f'{author_name}')
            author_run.bold = True
            
            title = book.title_and_statement_of_responsibility.title if book.title_and_statement_of_responsibility else ""
            
            # Обработка списка альтернативных имен
            alt_names = []
            if hasattr(book, 'personal_name_alternative'):
                alt_names = [name.entry_element for name in book.personal_name_alternative.all() if name and name.entry_element]
            alt_name = ", ".join(alt_names) if alt_names else ""
            
            # Обработка списка вторичных имен
            secondary_names = []
            if hasattr(book, 'personal_name_secondary'):
                secondary_names = [name.entry_element for name in book.personal_name_secondary.all() if name and name.entry_element]
            secondary_name = ", ".join(secondary_names) if secondary_names else ""
            
            place = book.publication_distribution.place_of_publication if book.publication_distribution else ""
            date = book.publication_distribution.date_of_publication if book.publication_distribution else ""
            dimensions = book.physical_description.dimensions if book.physical_description else ""
            note = book.summary_or_abstract.summary_text if book.summary_or_abstract else ""
            
            paragraph.add_run(f'{title}/'
                            f'{alt_name} '
                            f'{secondary_name}.'
                            f' - {place}:'
                            f'{alt_name}, '
                            f'{date}.'
                            f' - {dimensions}.\n'
                            f' - {note}')
            
            # Добавляем пустую строку между записями
            doc.add_paragraph()
        
        return doc
        
    except Exception as e:
        logger.error(f"Ошибка при генерации бюллетеня: {str(e)}")
        raise

def save_bulletin(doc, filename):
    """
    Сохраняет сгенерированный бюллетень в файл
    
    Args:
        doc: объект Document из python-docx
        filename: имя файла для сохранения
    """
    try:
        doc.save(filename)
    except Exception as e:
        logger.error(f"Ошибка при сохранении бюллетеня: {str(e)}")
        raise
