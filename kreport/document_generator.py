import calendar
import io
from collections import defaultdict
from datetime import datetime, timedelta

from django.db.models import Sum
from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
import tempfile
import logging

from klib.models import BaseOrder, BaseFundElement, BaseEdition, BaseOrderEdition, BaseArrival
from kservice.models import BookCirculation, EditionElement
from kuser.constants import EDUCATION_TYPE
from kuser.models import Reader

logger = logging.getLogger('main')


def generate_inventory_document(inventory_book):
    template_path = "kreport/DocFile/Инвентарная книга.docx"
    context = build_context_document(inventory_book)
    return create_docx(template_path, context)


def build_context_document(inventory_book):
    context = {
        'date_of_registration': inventory_book.date_of_registration.strftime("%d.%m.%Y"),
        'inventary_number': inventory_book.inventary_number,
        'year_of_publication': inventory_book.year_of_publication,
        'name': inventory_book.name,
        'author': inventory_book.author,
        'status': inventory_book.status,
        'document': inventory_book.document,
        'invoice_number': inventory_book.invoice_number,
        'invoice_date': (
            datetime.strptime(inventory_book.invoice_date, "%Y-%m-%d").strftime("%d.%m.%Y")
            if isinstance(inventory_book.invoice_date, str) else inventory_book.invoice_date.strftime("%d.%m.%Y")
        ),
        'price_without_vat': inventory_book.price_without_vat,
        'note': inventory_book.note,
    }
    return context

# def generate_summary_book_document(inventory_book):
#     template_path = "kreport/DocFile/Суммарная книга.docx"
#     context = build_context_document(inventory_book)
#     return create_docx(template_path, context)


def build_context_document(documents):
    context = {'records': documents}
    return context


def create_docx(template: str, contexts: list):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей
        doc.render({'records': contexts})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Failed to create file docx: {}'.format(e)
        return message


def generate_work_statistics(work_statistics):
    template_path = "kreport/DocFile/Распределение по должностям.docx"
    statistics = []
    for stat in work_statistics:
        context = build_context_work_statistics(stat)
        statistics.append(context)
    return create_docx_work_statistics(template_path, statistics)


def build_context_work_statistics(work_statistics):
    context = {
        'organization': work_statistics.organization,
        'department': work_statistics.department,
        'position': work_statistics.position,
        'user': work_statistics.user,
        'fio': f'{work_statistics.middle_name} {work_statistics.user.first_name} {work_statistics.user.last_name}',
        'date_of_registration': work_statistics.registration_date.strftime("%d.%m.%Y"),
        'status': work_statistics.work_type,
        'notes': work_statistics.notes,
    }
    return context


def generate_report_document(start_date, end_date, editions):
    template_path = "kreport/DocFile/Книга недополученных изданий.docx"

    statistics = []
    for edition in editions:
        number_of_copies = BaseOrder.objects.filter(
            edition_id=edition.id
        ).aggregate(total_copies=Sum('number_of_copies'))['total_copies'] or 0

        received_copies = BaseFundElement.objects.filter(
            edition=edition,
            registration_date__gte=start_date,
            registration_date__lt=end_date
        ).count()

        missing_copies = max(0, number_of_copies - received_copies)

        if missing_copies == 0 and received_copies == 0:
            continue

        base_order = BaseOrder.objects.filter(edition=edition).first()

        if base_order is None:
            company = "-"
            contract_number = "-"
        else:
            company = base_order.company if base_order.company else "-"
            contract_number = base_order.contract_number if base_order.contract_number else "-"

        context = build_context_book_incomplete_edition(edition, missing_copies, company, contract_number)
        statistics.append(context)

    return create_docx(template_path, statistics)


def build_context_book_incomplete_edition(editions, missing_copies, company, contract_number):
    context = {
        'title': editions.title,
        'missing_copies': missing_copies,
        'company': company,
        'contract_number': contract_number,
    }

    return context


def create_docx_work_statistics(template: str, contexts: list):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей
        doc.render({'records': contexts})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Failed to create file docx: {}'.format(e)
        return message


def get_education_name(education):
    for key, value in EDUCATION_TYPE:
        if key == education:
            return value
    return None


def generate_edu_statistics(edu_statistics):
    template_path = "kreport/DocFile/Распределение по образованию.docx"
    statistics = []
    for stat in edu_statistics:
        context = build_context_edu_statistics(stat)
        statistics.append(context)
    return create_docx_edu_statistics(template_path, statistics)


def build_context_edu_statistics(edu_statistics):
    context = {
        'user': edu_statistics.user,
        'fio': f'{edu_statistics.middle_name} {edu_statistics.user.first_name} {edu_statistics.user.last_name}',
        'organization': edu_statistics.organization,
        'position': edu_statistics.position,
        'education': get_education_name(edu_statistics.education),
    }
    return context


def create_docx_edu_statistics(template: str, contexts: list):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей
        doc.render({'records': contexts})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Failed to create file docx: {}'.format(e)
        return message


def create_total_book_report_file(document_data):
    doc = Document()

    doc.add_heading('Отчёт по книгам', level=1)

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Тип документа'
    hdr_cells[1].text = 'На начало периода'
    hdr_cells[2].text = 'Поступило'
    hdr_cells[3].text = 'На конец периода'

    for doc_data in document_data:
        row_cells = table.add_row().cells
        row_cells[0].text = doc_data['doc_name']
        row_cells[1].text = str(doc_data['doc_start'])
        row_cells[2].text = str(doc_data['doc_received'])
        row_cells[3].text = str(doc_data['doc_end'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


MONTHS = [
    (0, 'Весь год'),
    (1, 'Январь'),
    (2, 'Февраль'),
    (3, 'Март'),
    (4, 'Апрель'),
    (5, 'Май'),
    (6, 'Июнь'),
    (7, 'Июль'),
    (8, 'Август'),
    (9, 'Сентябрь'),
    (10, 'Октябрь'),
    (11, 'Ноябрь'),
    (12, 'Декабрь'),
]

def get_month_name(month_number):
    month_dict = dict(MONTHS)
    return month_dict.get(int(month_number))

def generate_user_statistics(user_statistics, month):
    template_path = "kreport/DocFile/Статистика месяц.docx"

    month_name = get_month_name(month)

    logger.debug(f'{user_statistics}')

    statistics = []
    for stat in user_statistics:
        context = build_context_user_statistics(stat, month)
        statistics.append(context)
    logger.debug(f'{statistics}')
    return create_docx_user_statistics(template_path, statistics, month_name)


def build_context_user_statistics(user_statistics, month):
    context = {
        'user': user_statistics.user,
        'fio': f'{user_statistics.middle_name} {user_statistics.user.first_name} {user_statistics.user.last_name}',
        'position': user_statistics.position,
        'date_of_registration': user_statistics.registration_date.strftime("%d.%m.%Y"),
        'month': month,
    }
    return context


def generate_reregistration_statistics(reader_data, month):
    template_path = "kreport/DocFile/Статистика месяц.docx"
    month_name = get_month_name(month)

    statistics = []
    for reader in reader_data:
        for rereg_date in reader.reregistration_dates.split():
            rereg_date_obj = datetime.strptime(rereg_date, "%Y-%m-%d").date()
            if month == 0 or rereg_date_obj.month == month:
                context = build_context_reregistration_statistics(reader, rereg_date_obj)
                statistics.append(context)

    logger.debug(f'{statistics}')
    return create_docx_user_statistics(template_path, statistics, month_name)


def build_context_reregistration_statistics(reader, reregistration_date):
    context = {
        'user': reader.user,
        'fio': f'{reader.middle_name} {reader.user.first_name} {reader.user.last_name}',
        'position': reader.position,
        'date_of_registration': reregistration_date.strftime("%d.%m.%Y"),
    }
    return context


def create_docx_user_statistics(template: str, contexts: list, month_name):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей, сгруппированных по месяцам
        doc.render({'records': contexts, 'month': month_name})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Не удалось создать файл docx: {}'.format(e)
        return message


def generate_debt_statistics(debt_statistics):
    template_path = "kreport/DocFile/Список должников.docx"
    statistics = []
    for stat in debt_statistics:
        context = build_context_debt_statistics(stat)
        statistics.append(context)
    return create_docx_debt_statistics(template_path, statistics)


def build_context_debt_statistics(debt_statistics):
    reader = Reader.objects.get(pk=debt_statistics.reader_id)
    context = {
        'fio': f'{reader.user.last_name} {reader.user.first_name} {reader.middle_name}',
        'inv_number': debt_statistics.inv_number,
        'service_type': debt_statistics.service_type,
        'refund_date': debt_statistics.refund_date.strftime("%d.%m.%Y"),
        'delay_days': debt_statistics.delay_days,
        }
    return context


def generate_debt_statistics_user(debt_statistics):
    template_path = "kreport/DocFile/Долги.docx"
    statistics = []
    for stat in debt_statistics:
        context = build_context_debt_user_statistics(stat)
        statistics.append(context)
    return create_docx_debt_statistics(template_path, statistics)


def build_context_debt_user_statistics(debt_statistics):
    circ = BookCirculation.objects.using('belrw-service-db').filter(pk=debt_statistics.circulation_id).first()
    context = {
        'inv_number': debt_statistics.inv_number,
        'service_type': debt_statistics.service_type,
        'status': circ.status,
        'refund_date': debt_statistics.refund_date.strftime("%d.%m.%Y"),
        'delay_days': debt_statistics.delay_days,
        }
    return context


def create_docx_debt_statistics(template: str, contexts: list):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей
        doc.render({'records': contexts})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Failed to create file docx: {}'.format(e)
        return message


def generate_report_document_non_per(start_date, end_date, editions):
    doc = Document()
    # Заголовок
    doc.add_heading('Отчет о недополученных изданиях', level=1)
    doc.add_paragraph(f'Период: {start_date.strftime("%d.%m.%Y")} - {end_date.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Сформировано: {datetime.now().strftime("%d.%m.%Y %H:%M")}', style='Normal')

    # Таблица
    table = doc.add_table(rows=1, cols=5, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Тип издания'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Недостающие экземпляры'
    hdr_cells[3].text = 'Организация'
    hdr_cells[4].text = 'Договор'

    for edition in editions:
        total_ordered_copies = BaseOrderEdition.objects.filter(edition=edition).aggregate(
            total_quantity=Sum('quantity')
        )['total_quantity'] or 0

        received_copies = BaseArrival.objects.filter(
            edition=edition,
            invoice_date__gte=start_date,
            invoice_date__lt=end_date
        ).aggregate(total_qty=Sum('qty'))['total_qty'] or 0

        missing_copies = max(0, total_ordered_copies - received_copies)

        if missing_copies == 0:
            continue

        base_order = BaseOrder.objects.filter(edition=edition).first()

        if base_order is None:
            company = "-"
            contract_number = "-"
        else:
            company = base_order.company if base_order.company else "-"
            contract_number = base_order.contract_number if base_order.contract_number else "-"

        row_cells = table.add_row().cells
        row_cells[0].text = dict(BaseEdition.SUBTYPES).get(edition.edition_subtype, "N/A")
        row_cells[1].text = edition.title
        row_cells[2].text = str(missing_copies)
        row_cells[3].text = str(company)
        row_cells[4].text = contract_number

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

    doc.add_paragraph(f'\nИтог: Всего недостающих экземпляров — {total_missing_copies}', style='Bold')

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def generate_accountiong_statement_document(accounting_statement):
    document = Document()
    document.add_heading('Отчет по бухгалтерскому учету', level=1)
    document.add_paragraph(f"Год: {accounting_statement.year}")
    document.add_paragraph(f"Месяц: {accounting_statement.month}")

    year = int(accounting_statement.year)
    month = int(accounting_statement.month)

    arrivals = BaseArrival.objects.filter(
        invoice_date__year=year,
        invoice_date__month=month
    ).order_by('order_edition__order__company')

    grouped_data = {}
    for arrival in arrivals:
        company = (
            arrival.order_edition.order.company
            if arrival.order_edition and arrival.order_edition.order
            else "Без компании"
        )
        if company not in grouped_data:
            grouped_data[company] = []
        grouped_data[company].append(arrival)

    for company, company_arrivals in grouped_data.items():
        document.add_heading(f'Компания: {company}', level=2)

        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Тип сопр. документа'
        header_cells[1].text = 'Номер документа'
        header_cells[2].text = 'Дата документа'
        header_cells[3].text = 'Сумма с НДС'
        header_cells[4].text = 'На балансе (Кол-во/Сумма/НДС)'
        header_cells[5].text = 'Не на балансе (Кол-во/Сумма/НДС)'

        company_total_amount_with_vat = 0
        company_total_balance_qty = 0
        company_total_balance_amount = 0
        company_total_balance_vat = 0

        for arrival in company_arrivals:
            row_cells = table.add_row().cells
            row_cells[0].text = arrival.filing or ''
            row_cells[1].text = arrival.invoice_number or ''
            row_cells[2].text = str(arrival.invoice_date) or ''
            row_cells[3].text = str(arrival.amount_with_vat or 0)

            balance_qty = arrival.qty if arrival.balance_type == BaseArrival.BALANCE_TYPE_1.lower() else 0
            balance_amount = arrival.amount if arrival.balance_type == BaseArrival.BALANCE_TYPE_1.lower() else 0
            balance_vat = arrival.amount_vat if arrival.balance_type == BaseArrival.BALANCE_TYPE_1.lower() else 0

            no_balance_qty = arrival.qty if arrival.balance_type == BaseArrival.BALANCE_TYPE_2.lower() else 0
            no_balance_amount = arrival.amount if arrival.balance_type == BaseArrival.BALANCE_TYPE_2.lower() else 0
            no_balance_vat = arrival.amount_vat if arrival.balance_type == BaseArrival.BALANCE_TYPE_2.lower() else 0

            row_cells[4].text = f"{balance_qty} / {balance_amount} / {balance_vat}"
            row_cells[5].text = f"{no_balance_qty} / {no_balance_amount} / {no_balance_vat}"

            company_total_amount_with_vat += arrival.amount_with_vat or 0

            if arrival.balance_type == BaseArrival.BALANCE_TYPE_1.lower():
                company_total_balance_qty += balance_qty
                company_total_balance_amount += balance_amount
                company_total_balance_vat += balance_vat

        row_cells = table.add_row().cells
        row_cells[0].text = f"Итого по {company}"
        row_cells[3].text = str(company_total_amount_with_vat)
        row_cells[4].text = f"{company_total_balance_qty} / {company_total_balance_amount} / {company_total_balance_vat}"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_certificate_acceptance_document(certificate):
    document = Document()
    document.add_heading('Отчет по приему периодических изданий', level=1)
    document.add_paragraph(f"Год: {certificate.year}")
    year = int(certificate.year)
    editions = BaseArrival.objects.filter(invoice_date__year=year)
    grouped_data = {}
    for edition in editions:
        name = edition.edition.title if edition.edition else "N/A"
        if name not in grouped_data:
            grouped_data[name] = {
                "years": set(),
                "numbers": {},
                "double_numbers": [],
                "total_qty": 0
            }
        grouped_data[name]["years"].add(edition.invoice_date.year if edition.invoice_date else "N/A")
        invoice_number = edition.invoice_number if edition.invoice_number else "N/A"
        grouped_data[name]["numbers"].setdefault(invoice_number, 0)
        grouped_data[name]["numbers"][invoice_number] += edition.qty or 0
        grouped_data[name]["double_numbers"].append(
            "Да" if edition.double_number and edition.double_number.is_double else "Нет")
        grouped_data[name]["total_qty"] += edition.qty or 0

    grand_total = 0

    for name, data in grouped_data.items():
        document.add_heading(f'Наименование: {name}', level=2)

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Год'
        header_cells[1].text = 'Номер документа'
        header_cells[2].text = 'Сдвоенный'
        header_cells[3].text = 'Количество'
        header_cells[4].text = 'Итого экземпляров'

        years = ', '.join(map(str, data["years"]))
        invoice_details = ', '.join(data["numbers"].keys())
        double_numbers = ', '.join(data["double_numbers"])
        quantities = ', '.join(map(str, data["numbers"].values()))
        total_copies = data["total_qty"]

        row_cells = table.add_row().cells
        row_cells[0].text = years
        row_cells[1].text = invoice_details
        row_cells[2].text = double_numbers
        row_cells[3].text = quantities
        row_cells[4].text = str(total_copies)

        grand_total += total_copies

    document.add_paragraph(f'Общий итог экземпляров: {grand_total}', style='Intense Quote')

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_periodicals_acceptance_document(report):
    document = Document()
    document.add_heading('Отчет о приеме периодических изданий', level=1)
    document.add_paragraph(f"Год: {report.year}, Месяц: {report.month}")

    year = int(report.year)
    month = int(report.month)
    arrivals = BaseArrival.objects.filter(
        invoice_date__year=year,
        invoice_date__month=month
    ).select_related('edition', 'order_edition__order__company')

    grouped_data = {}
    for arrival in arrivals:
        company = arrival.order_edition.order.company
        edition_name = arrival.edition.title if arrival.edition else "—"
        qty = arrival.qty
        amount_no_vat = arrival.amount
        vat_amount = arrival.amount_vat
        contract_number = arrival.order_edition.order.contract_number or "—"
        contract_date = arrival.order_edition.order.contract_date.strftime(
            '%d.%m.%Y') if arrival.order_edition.order.contract_date else "—"

        if company not in grouped_data:
            grouped_data[company] = {}

        if edition_name not in grouped_data[company]:
            grouped_data[company][edition_name] = {
                "arrivals": [],
                "total_qty": 0,
                "total_amount_no_vat": 0,
                "total_vat": 0
            }

        grouped_data[company][edition_name]["arrivals"].append({
            "qty": qty,
            "amount_no_vat": amount_no_vat,
            "vat_amount": vat_amount,
            "contract_number": contract_number,
            "contract_date": contract_date
        })
        grouped_data[company][edition_name]["total_qty"] += qty
        grouped_data[company][edition_name]["total_amount_no_vat"] += amount_no_vat
        grouped_data[company][edition_name]["total_vat"] += vat_amount

    grand_total_qty = 0
    grand_total_amount_no_vat = 0
    grand_total_vat = 0

    for company, editions in grouped_data.items():
        document.add_heading(f'Организация: {company}', level=2)

        for edition_name, data in editions.items():
            para = document.add_paragraph(f'Наименование издания: {edition_name}', style='Normal')
            table = document.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            header_cells = table.rows[0].cells
            header_cells[0].text = 'Кол-во'
            header_cells[1].text = 'Цена (без НДС)'
            header_cells[2].text = 'НДС'
            header_cells[3].text = 'Номер договора'
            header_cells[4].text = 'Дата договора'

            for arrival in data["arrivals"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(arrival["qty"])
                row_cells[1].text = f'{arrival["amount_no_vat"]:.2f}'
                row_cells[2].text = f'{arrival["vat_amount"]:.2f}'
                row_cells[3].text = arrival["contract_number"]
                row_cells[4].text = arrival["contract_date"]

            summary_text = f'Итого по изданию "{edition_name}": ' \
                           f'Количество: {data["total_qty"]}, ' \
                           f'Сумма без НДС: {data["total_amount_no_vat"]:.2f}, ' \
                           f'Сумма НДС: {data["total_vat"]:.2f}'
            document.add_paragraph(summary_text, style='Normal')

            grand_total_qty += data["total_qty"]
            grand_total_amount_no_vat += data["total_amount_no_vat"]
            grand_total_vat += data["total_vat"]

    final_summary = f'Общий итог по всем организациям: ' \
                    f'Количество: {grand_total_qty}, ' \
                    f'Сумма без НДС: {grand_total_amount_no_vat:.2f}, ' \
                    f'Сумма НДС: {grand_total_vat:.2f}'
    document.add_paragraph(final_summary, style='Normal')

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_circulation_statistics(debt_statistics):
    template_path = "kreport/DocFile/Статистика по обороту(год).docx"
    statistics = []
    for stat in debt_statistics:
        context = build_context_circulation_statistics(stat)
        statistics.append(context)
    return create_docx_circulation_statistics(template_path, statistics)


def generate_circulation_statistics_user(debt_statistics):
    template_path = "kreport/DocFile/Статистика по обороту(пользователь).docx"
    statistics = []
    for stat in debt_statistics:
        context = build_context_circulation_statistics_user(stat)
        statistics.append(context)
    return create_docx_circulation_statistics(template_path, statistics)


def build_context_circulation_statistics(circulation_statistics):
    reader = Reader.objects.get(pk=circulation_statistics.reader_id)
    edition_element = EditionElement.objects.using('belrw-service-db').filter(
        book_circulation=circulation_statistics).first()
    fund = BaseFundElement.objects.using('belrw-lib-db').filter(pk=edition_element.fund_id).first()
    edition_type = ''
    if fund.inventory_number.split('/')[0] == '1':
        edition_type = 'Книги'
    if fund.inventory_number.split('/')[0] == '2':
        edition_type = 'НТД и брошюры'
    if fund.inventory_number.split('/')[0] == '3':
        edition_type = 'Журналы и другие периодические издания'
    if fund.inventory_number.split('/')[0] == '4':
        edition_type = 'Газеты'
    if fund.inventory_number.split('/')[0] == '5':
        edition_type = 'Информационные листки'
    if fund.inventory_number.split('/')[0] == '6':
        edition_type = 'Изменения и дополнения к НТД'
    if fund.inventory_number.split('/')[0] == '8':
        edition_type = 'Электронный ресурс'

    context = {
        'fio': f'{reader.user.last_name} {reader.user.first_name} {reader.middle_name}',
        'inv_number': fund.inventory_number,
        'service_type': circulation_statistics.service_type,
        'refund_date': circulation_statistics.refund_date.strftime("%d.%m.%Y"),
        'receive_date': circulation_statistics.receive_date.strftime("%d.%m.%Y"),
        'edition_type': edition_type,
        'edition': edition_element.edition,
        'status': circulation_statistics.status,
        }
    return context


def build_context_circulation_statistics_user(circulation_statistics):
    edition_element = EditionElement.objects.using('belrw-service-db').filter(
        book_circulation=circulation_statistics).first()
    fund = BaseFundElement.objects.using('belrw-lib-db').filter(pk=edition_element.fund_id).first()

    edition_type = ''
    if fund.inventory_number.split('/')[0] == '1':
        edition_type = 'Книги'
    if fund.inventory_number.split('/')[0] == '2':
        edition_type = 'НТД и брошюры'
    if fund.inventory_number.split('/')[0] == '3':
        edition_type = 'Журналы и другие периодические издания'
    if fund.inventory_number.split('/')[0] == '4':
        edition_type = 'Газеты'
    if fund.inventory_number.split('/')[0] == '5':
        edition_type = 'Информационные листки'
    if fund.inventory_number.split('/')[0] == '6':
        edition_type = 'Изменения и дополнения к НТД'
    if fund.inventory_number.split('/')[0] == '8':
        edition_type = 'Электронный ресурс'

    context = {
        'inv_number':  fund.inventory_number,
        'service_type': circulation_statistics.service_type,
        'refund_date': circulation_statistics.refund_date.strftime("%d.%m.%Y"),
        'receive_date': circulation_statistics.receive_date.strftime("%d.%m.%Y"),
        'edition_type': edition_type,
        'edition': edition_element.edition,
        'status': circulation_statistics.status,
        }
    return context


def create_docx_circulation_statistics(template: str, contexts: list):
    try:
        doc = DocxTemplate(template)
        # Рендерим документ с данными всех записей
        doc.render({'records': contexts})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

        doc.save(temp_file.name)
        return open(temp_file.name, 'rb')
    except Exception as e:
        message = 'Failed to create file docx: {}'.format(e)
        return message


def generate_typeofservice_document(report):
    document = Document()
    document.add_heading('Отчет по типам обслуживания', level=1)

    year = int(report.year)
    month = int(report.month)

    start_date = datetime(year, month, 1)
    if month == 12:
        end_date = datetime(year + 1, 1, 1)
    else:
        end_date = datetime(year, month + 1, 1)

    circulations = BookCirculation.objects.filter(
        receive_date__gte=start_date,
    )

    service_type_count = {}

    for circulation in circulations:
        service_type = circulation.service_type
        if service_type in service_type_count:
            service_type_count[service_type] += 1
        else:
            service_type_count[service_type] = 1

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Тип обслуживания'
    hdr_cells[1].text = 'Количество'

    for service_type, count in service_type_count.items():
        row_cells = table.add_row().cells
        row_cells[0].text = service_type
        row_cells[1].text = str(count)

    total_count = sum(service_type_count.values())
    document.add_paragraph(f"\nИтоговое количество обслуживаний: {total_count}")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def generate_storage_report_document(report):
    document = Document()
    document.add_heading('Отчет по подтипам изданий', level=1)

    year = int(report.year)
    month = int(report.month)

    start_date = datetime(year, month, 1)
    if month == 12:
        end_date = datetime(year + 1, 1, 1)
    else:
        end_date = datetime(year, month + 1, 1)

    edition_elements = EditionElement.objects.filter(
        book_circulation__receive_date__gte=start_date,
        book_circulation__receive_date__lt=end_date
    ).order_by('book_circulation__receive_date')

    subtype_dict = dict(BaseEdition.SUBTYPES)
    subtype_count = {}

    for edition_element in edition_elements:
        base_edition = BaseEdition.objects.filter(title=edition_element.edition).first()
        if base_edition:
            edition_subtype = subtype_dict.get(base_edition.edition_subtype, 'Неизвестно')
        else:
            edition_subtype = 'Неизвестно'

        if edition_subtype in subtype_count:
            subtype_count[edition_subtype] += 1
        else:
            subtype_count[edition_subtype] = 1

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Подтип издания'
    hdr_cells[1].text = 'Количество'

    for subtype, count in subtype_count.items():
        row_cells = table.add_row().cells
        row_cells[0].text = subtype
        row_cells[1].text = str(count)

    total_count = sum(subtype_count.values())
    document.add_paragraph(f"\nИтоговое количество изданий: {total_count}")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
