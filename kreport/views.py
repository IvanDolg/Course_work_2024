import logging
from datetime import date, datetime, timedelta

import json
from django.http import JsonResponse, HttpResponse, Http404
from docx import Document
from django.db.models import Sum

from kservice.models import Debtors, BookCirculation
from kuser.models import Reader
from .constans import FIRST_BOOK_TYPE, SECOND_BOOK_TYPE, \
    THEAD_BOOK_TYPE
from klib.models import BaseFundElement, BaseArrival, BaseEdition, BaseOrder, Edition, BaseOrderEdition
from .document_calculations import calculate_doc_start_total_document, calculate_doc_received_total_document, \
    calculate_doc_end_document, calculate_doc_start_accepted_to_balance, calculate_doc_received_accepted_to_balance, \
    calculate_doc_end_accepted_to_balance, calculate_doc_start_books, calculate_doc_received_books, \
    calculate_doc_end_books, calculate_doc_start_electronic_resources, calculate_doc_received_electronic_resources, \
    calculate_doc_end_electronic_resources, \
    calculate_doc_start_brochures, calculate_doc_received_brochures, calculate_doc_end_brochures, \
    calculate_doc_start_ntd, calculate_doc_received_ntd, calculate_doc_end_ntd, calculate_doc_start_information_sheets, \
    calculate_doc_received_information_sheets, calculate_doc_end_information_sheets, calculate_doc_start_magazines, \
    calculate_doc_received_magazines, calculate_doc_end_magazines, calculate_doc_end_newspapers, \
    calculate_doc_received_newspapers, calculate_doc_start_newspapers
from .document_generator import build_context_document, create_docx, \
    generate_work_statistics, generate_edu_statistics, generate_report_document, \
    create_total_book_report_file, generate_user_statistics, generate_reregistration_statistics, \
    generate_debt_statistics, generate_debt_statistics_user, generate_report_document_non_per, \
    generate_accountiong_statement_document, generate_certificate_acceptance_document, \
    generate_periodicals_acceptance_document, generate_typeofservice_document, generate_circulation_statistics, \
    generate_circulation_statistics_user, generate_storage_report_document
from .models import CreateWorkplaceReport, CreateEducationReport, BookIncompleteEdition, TotalBook, CreateInventoryBook, \
    UserAccounting, CreateDebtorsReport, BookIncompleteEditionNonPeriodicals, AccountiongStatement, \
    CertificateAcceptancePeriodicals, PeriodicalsAcceptanceReport, TypeOfService, StorageReport, BookCirculationReport

logger = logging.getLogger('main')


def inventory_book_report(request, object_id):
    if request.method == 'POST':
        try:
            obj = CreateInventoryBook.objects.get(pk=object_id)
        except CreateInventoryBook.DoesNotExist:
            raise Http404("CreateInventoryBook object does not exist")

        # Получение необходимых данных
        first_inventory_number = BaseFundElement.objects.filter(pk=obj.first_inventory_number).first()
        last_inventory_number = BaseFundElement.objects.filter(pk=obj.last_inventory_number).first()

        if not first_inventory_number or not last_inventory_number:
            return JsonResponse({"error": "Не указаны начальный или конечный инвентарные номера."}, status=400)

        try:
            first_prefix, first_suffix = map(int, first_inventory_number.inventory_number.split('/'))
            last_prefix, last_suffix = map(int, last_inventory_number.inventory_number.split('/'))
        except ValueError:
            return JsonResponse({"error": "Некорректный формат инвентарных номеров."}, status=400)

        if first_prefix != last_prefix:
            return JsonResponse({"error": "Префиксы начального и конечного номеров не совпадают."}, status=400)

        publication_status = None
        if obj.display_excluded_editions:
            publication_status = BaseFundElement.PUBLICATION_STATUS_WRITTEN_OFF
        elif obj.display_current_editions:
            publication_status = BaseFundElement.PUBLICATION_STATUS_NOT_WRITTEN_OFF

        fund_elements = BaseFundElement.objects.filter(
            inventory_number__startswith=f"{first_prefix}/",
        )

        if publication_status is not None:
            fund_elements = fund_elements.filter(publication_status=publication_status)

        if obj.template_type:
            fund_elements = fund_elements.filter(edition__edition_subtype=obj.template_type)

        fund_elements_in_range = [
            el for el in fund_elements
            if first_suffix <= int(el.inventory_number.split('/')[1]) <= last_suffix
        ]

        # Создаем документ Word
        document = Document()
        document.add_heading('Отчет по инвентарной книге', level=1)
        document.add_paragraph(f'Дата создания отчета: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')

        # Создаем таблицу
        table = document.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = [
            "Дата регистрации", "Инвентарные номера", "Год публикации",
            "Название", "Автор", "Статус", "Номер накладной",
            "Дата накладной", "Цена без НДС"
        ]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        for el in fund_elements_in_range:
            edition = el.get_edition()
            if not edition:
                continue  # Skip if there's no edition

            row_cells = table.add_row().cells
            row_cells[0].text = edition.created_at.strftime('%Y-%m-%d') if edition.created_at else 'N/A'
            row_cells[1].text = el.inventory_number
            row_cells[2].text = str(edition.year) if edition.year else 'N/A'
            row_cells[3].text = edition.title if edition.title else 'N/A'
            row_cells[4].text = edition.author if edition.author else 'N/A'
            row_cells[5].text = el.get_publication_status_display() if el.get_publication_status_display() else 'N/A'
            row_cells[6].text = el.invoice_number or 'N/A'
            row_cells[7].text = el.invoice_date.strftime('%Y-%m-%d') if el.invoice_date else 'N/A'
            row_cells[8].text = str(el.price) if el.price is not None else 'N/A'

        doc_file = f'/tmp/inventory_report_{object_id}.docx'
        document.save(doc_file)

        with open(doc_file, 'rb') as f:
            response = HttpResponse(
                f.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="InventoryBookReport_{object_id}.docx"'
            return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def get_inventory_numbers(request):
    template_type = request.GET.get('template_type')
    inventory_numbers = BaseFundElement.objects.filter(edition__edition_subtype=template_type).values('id',
                                                                                                      'inventory_number')
    return JsonResponse({'inventory_numbers': list(inventory_numbers)})


def generate_summary_book_document(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        documents = data.get('documents', [])
        context = build_context_document(documents)
        template_path = "kreport/DocFile/Суммарная книга.docx"
        doc_file = create_docx(template_path, context)
        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Суммарная_книга.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_workplace_report(request, object_id):
    if request.method == 'POST':
        obj: CreateWorkplaceReport = CreateWorkplaceReport.objects.get(pk=object_id)

        readers = Reader.objects.using('belrw-user-db').all()
        user = Reader.objects.using('belrw-user-db').get(user=request.user)

        if user.library is not None:
            readers = readers.filter(library=user.library)

        if obj.organization is not None:
            readers = readers.filter(organization__name=obj.organization)

        if obj.position is not None:
            readers = readers.filter(position__name=obj.position)

        if obj.year is not None:
            readers = readers.filter(registration_date__year=obj.year)

        if not obj.add_excluded:
            readers = readers.exclude(exclusion=True)

        doc_file = generate_work_statistics(readers)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Распределение_по_должностям.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_workplace_report_first_class(request, object_id):
    if request.method == 'POST':
        obj: CreateWorkplaceReport = CreateWorkplaceReport.objects.get(pk=object_id)
        readers = Reader.objects.using('belrw-user-db').all()

        if obj.library is not None:
            readers = readers.filter(library=obj.library)

        if obj.organization is not None:
            readers = readers.filter(organization__name=obj.organization)

        if obj.position is not None:
            readers = readers.filter(position__name=obj.position)

        if obj.year is not None:
            readers = readers.filter(registration_date__year=obj.year)

        if not obj.add_excluded:
            readers = readers.exclude(exclusion=True)

        doc_file = generate_work_statistics(readers)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Распределение_по_должностям.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


# BookIncompleteEdition
def generate_book_incomplete_edition_report(request, object_id):
    if request.method == 'POST':
        try:
            obj = BookIncompleteEdition.objects.get(pk=object_id)
        except BookIncompleteEdition.DoesNotExist:
            raise Http404("BookIncompleteEdition object does not exist")

        start_date = datetime(obj.year, obj.moth, 1)
        end_date = (start_date + timedelta(days=31)).replace(day=1)
        editions = BaseEdition.objects.all()
        filtered_editions = []

        for edition in editions:
            number_of_copies = BaseOrder.objects.filter(
                edition_id=edition.id
            ).aggregate(total_copies=Sum('number_of_copies'))['total_copies'] or 0

            received_copies = BaseFundElement.objects.filter(
                edition=edition,
                registration_date__gte=start_date,
                registration_date__lt=end_date
            ).count()

            missing_copies = max(0, number_of_copies - received_copies)

            if missing_copies == 0 or received_copies == 0:
                continue

            filtered_editions.append(edition)
        doc_file = generate_report_document_non_per(start_date, end_date, filtered_editions)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="BookIncompleteEditionReport.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_book_incomplete_edition_report_non_per(request, object_id):
    if request.method == 'POST':
        try:
            obj = BookIncompleteEditionNonPeriodicals.objects.get(pk=int(object_id))
        except BookIncompleteEditionNonPeriodicals.DoesNotExist:
            raise Http404("BookIncompleteEditionNonPeriodicals object does not exist")
        except ValueError:
            raise Http404("Invalid object ID")

        year = int(obj.year)
        month = int(obj.month)

        start_date = datetime(year, month, 1)
        end_date = (start_date + timedelta(days=31)).replace(day=1)
        editions = BaseEdition.objects.filter(edition_type=Edition.TYPE_NON_PERIODICAL)
        filtered_editions = []

        for edition in editions:
            total_ordered_copies = BaseOrderEdition.objects.filter(edition=edition).aggregate(
                total_quantity=Sum('quantity')
            )['total_quantity'] or 0

            received_copies = BaseArrival.objects.filter(
                edition=edition,
                invoice_date__gte=start_date,
                invoice_date__lt=end_date
            ).aggregate(total_qty=Sum('qty'))['total_qty'] or 0

            missing_copies = max(0, total_ordered_copies - received_copies)

            if missing_copies == 0:
                continue

            filtered_editions.append(edition)
        doc_file = generate_report_document_non_per(start_date, end_date, filtered_editions)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="BookIncompleteEditionReport.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)



def get_document_data(obj, year, month):
    if obj.book_type == FIRST_BOOK_TYPE:
        return [
            {"doc_name": "Всего документов", "doc_start": calculate_doc_start_total_document(year, month),
             "doc_received": calculate_doc_received_total_document(year, month),
             "doc_end": calculate_doc_end_document(year, month)},
            {"doc_name": "В.т.ч принятых на баланс",
             "doc_start": calculate_doc_start_accepted_to_balance(year, month),
             "doc_received": calculate_doc_received_accepted_to_balance(year, month),
             "doc_end": calculate_doc_end_accepted_to_balance(year, month)},
            {"doc_name": "Книги", "doc_start": calculate_doc_start_books(year, month),
             "doc_received": calculate_doc_received_books(year, month),
             "doc_end": calculate_doc_end_books(year, month)},
            {"doc_name": "Электронные ресурсы", "doc_start": calculate_doc_start_electronic_resources(year, month),
             "doc_received": calculate_doc_received_electronic_resources(year, month),
             "doc_end": calculate_doc_end_electronic_resources(year, month)},
            {"doc_name": "Брошюры", "doc_start": calculate_doc_start_brochures(year, month),
             "doc_received": calculate_doc_received_brochures(year, month),
             "doc_end": calculate_doc_end_brochures(year, month)},
            {"doc_name": "НТД", "doc_start": calculate_doc_start_ntd(year, month),
             "doc_received": calculate_doc_received_ntd(year, month),
             "doc_end": calculate_doc_end_ntd(year, month)},
            {"doc_name": "Информационные листки", "doc_start": calculate_doc_start_information_sheets(year, month),
             "doc_received": calculate_doc_received_information_sheets(year, month),
             "doc_end": calculate_doc_end_information_sheets(year, month)},
            {"doc_name": "Журналы", "doc_start": calculate_doc_start_magazines(year, month),
             "doc_received": calculate_doc_received_magazines(year, month),
             "doc_end": calculate_doc_end_magazines(year, month)},
            {"doc_name": "Газеты", "doc_start": calculate_doc_start_newspapers(year, month),
             "doc_received": calculate_doc_received_newspapers(year, month),
             "doc_end": calculate_doc_end_newspapers(year, month)},
        ]

    elif obj.book_type == SECOND_BOOK_TYPE:
        return [
            {"doc_name": "Всего документов", "doc_start": calculate_doc_start_total_document(year, month),
             "doc_received": calculate_doc_received_total_document(year, month),
             "doc_end": calculate_doc_end_document(year, month)},
            {"doc_name": "В.т.ч принятых на баланс",
             "doc_start": calculate_doc_start_accepted_to_balance(year, month),
             "doc_received": calculate_doc_received_accepted_to_balance(year, month),
             "doc_end": calculate_doc_end_accepted_to_balance(year, month)},
            {"doc_name": "Книги", "doc_start": calculate_doc_start_books(year, month),
             "doc_received": calculate_doc_received_books(year, month),
             "doc_end": calculate_doc_end_books(year, month)},
            {"doc_name": "Электронные ресурсы", "doc_start": calculate_doc_start_electronic_resources(year, month),
             "doc_received": calculate_doc_received_electronic_resources(year, month),
             "doc_end": calculate_doc_end_electronic_resources(year, month)},
            {"doc_name": "Брошюры", "doc_start": calculate_doc_start_brochures(year, month),
             "doc_received": calculate_doc_received_brochures(year, month),
             "doc_end": calculate_doc_end_brochures(year, month)},
            {"doc_name": "НТД", "doc_start": calculate_doc_start_ntd(year, month),
             "doc_received": calculate_doc_received_ntd(year, month),
             "doc_end": calculate_doc_end_ntd(year, month)},
            {"doc_name": "Информационные листки", "doc_start": calculate_doc_start_information_sheets(year, month),
             "doc_received": calculate_doc_received_information_sheets(year, month),
             "doc_end": calculate_doc_end_information_sheets(year, month)},
            {"doc_name": "Журналы", "doc_start": calculate_doc_start_magazines(year, month),
             "doc_received": calculate_doc_received_magazines(year, month),
             "doc_end": calculate_doc_end_magazines(year, month)},
            {"doc_name": "Газеты", "doc_start": calculate_doc_start_newspapers(year, month),
             "doc_received": calculate_doc_received_newspapers(year, month),
             "doc_end": calculate_doc_end_newspapers(year, month)},
        ]

    elif obj.book_type == THEAD_BOOK_TYPE:
        return [
            {"doc_name": "Всего документов", "doc_start": calculate_doc_start_total_document(year, month),
             "doc_received": calculate_doc_received_total_document(year, month),
             "doc_end": calculate_doc_end_document(year, month)},
            {"doc_name": "В.т.ч принятых на баланс",
             "doc_start": calculate_doc_start_accepted_to_balance(year, month),
             "doc_received": calculate_doc_received_accepted_to_balance(year, month),
             "doc_end": calculate_doc_end_accepted_to_balance(year, month)},
            {"doc_name": "Книги", "doc_start": calculate_doc_start_books(year, month),
             "doc_received": calculate_doc_received_books(year, month),
             "doc_end": calculate_doc_end_books(year, month)},
            {"doc_name": "Электронные ресурсы", "doc_start": calculate_doc_start_electronic_resources(year, month),
             "doc_received": calculate_doc_received_electronic_resources(year, month),
             "doc_end": calculate_doc_end_electronic_resources(year, month)},
            {"doc_name": "Брошюры", "doc_start": calculate_doc_start_brochures(year, month),
             "doc_received": calculate_doc_received_brochures(year, month),
             "doc_end": calculate_doc_end_brochures(year, month)},
            {"doc_name": "НТД", "doc_start": calculate_doc_start_ntd(year, month),
             "doc_received": calculate_doc_received_ntd(year, month),
             "doc_end": calculate_doc_end_ntd(year, month)},
            {"doc_name": "Информационные листки", "doc_start": calculate_doc_start_information_sheets(year, month),
             "doc_received": calculate_doc_received_information_sheets(year, month),
             "doc_end": calculate_doc_end_information_sheets(year, month)},
            {"doc_name": "Журналы", "doc_start": calculate_doc_start_magazines(year, month),
             "doc_received": calculate_doc_received_magazines(year, month),
             "doc_end": calculate_doc_end_magazines(year, month)},
            {"doc_name": "Газеты", "doc_start": calculate_doc_start_newspapers(year, month),
             "doc_received": calculate_doc_received_newspapers(year, month),
             "doc_end": calculate_doc_end_newspapers(year, month)},
        ]
    else:
        return []


def generate_total_book_report(request, object_id):
    if request.method == 'POST':
        obj = TotalBook.objects.get(pk=object_id)

        year = int(obj.year)
        month = int(obj.month)
        document_data = get_document_data(obj, year, month)

        doc_file = create_total_book_report_file(document_data)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Total_Book_Report.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_education_report(request, object_id):
    if request.method == 'POST':
        obj: CreateEducationReport = CreateEducationReport.objects.get(pk=object_id)
        readers = Reader.objects.using('belrw-user-db').all()
        user = Reader.objects.using('belrw-user-db').get(user=request.user)

        if user.library is not None:
            readers = readers.filter(library=user.library)

        if obj.education is not None:
            readers = readers.filter(education=obj.education)

        if obj.year is not None:
            readers = readers.filter(registration_date__year=obj.year)

        if not obj.add_excluded:
            readers = readers.exclude(exclusion=True)

        doc_file = generate_edu_statistics(readers)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Распределение по образованию.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_education_report_first_class(request, object_id):
    if request.method == 'POST':
        obj: CreateEducationReport = CreateEducationReport.objects.get(pk=object_id)
        readers = Reader.objects.using('belrw-user-db')

        if obj.library is not None:
            readers = readers.filter(library=obj.library)

        if obj.education is not None:
            readers = readers.filter(education=obj.education)

        if obj.year is not None:
            readers = readers.filter(registration_date__year=obj.year)

        if not obj.add_excluded:
            readers = readers.exclude(exclusion=True)

        doc_file = generate_edu_statistics(readers)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Распределение по образованию.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_user_accounting(request, object_id):
    if request.method == 'POST':
        obj: UserAccounting = UserAccounting.objects.get(pk=object_id)
        readers = Reader.objects.using('belrw-user-db').all()
        user = Reader.objects.using('belrw-user-db').get(user=request.user)

        if user.library is not None:
            readers = readers.filter(library=user.library)

        if obj.report_type == 'Registered':
            if obj.month is not None and obj.month != 0:
                readers = readers.filter(registration_date__month=obj.month)

            if obj.year is not None:
                readers = readers.filter(registration_date__year=obj.year)

            if obj.day is not None and obj.day != 0 and obj.month is not None and obj.month != 0:
                readers = readers.filter(registration_date__day=obj.day)

            logger.debug(f'readers: {readers}')

            doc_file = generate_user_statistics(readers, obj.month)
        else:
            reader_data = []
            for reader in readers:
                if reader.reregistration_dates:
                    if obj.year:
                        valid_dates = [
                            date for date in reader.reregistration_dates.split()
                            if datetime.strptime(date, "%Y-%m-%d").year == obj.year and
                               (obj.month == 0 or datetime.strptime(date, "%Y-%m-%d").month == obj.month)
                        ]
                        if valid_dates:
                            reader_data.append(reader)

            doc_file = generate_reregistration_statistics(reader_data, obj.month)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Просмотр статистики.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def generate_user_accounting_first_class(request, object_id):
    if request.method == 'POST':
        obj: UserAccounting = UserAccounting.objects.get(pk=object_id)
        readers = Reader.objects.using('belrw-user-db').all()

        if obj.library is not None:
            readers = readers.filter(library=obj.library)

        if obj.report_type == 'Registered':
            if obj.month is not None and obj.month != 0:
                readers = readers.filter(registration_date__month=obj.month)

            if obj.year is not None:
                readers = readers.filter(registration_date__year=obj.year)

            if obj.day is not None and obj.day != 0 and obj.month is not None and obj.month != 0:
                readers = readers.filter(registration_date__day=obj.day)

            logger.debug(f'readers: {readers}')

            doc_file = generate_user_statistics(readers, obj.month)
        else:
            reader_data = []
            for reader in readers:
                if reader.reregistration_dates:
                    if obj.year:
                        valid_dates = [
                            date for date in reader.reregistration_dates.split()
                            if datetime.strptime(date, "%Y-%m-%d").year == obj.year and
                               (obj.month == 0 or datetime.strptime(date, "%Y-%m-%d").month == obj.month)
                        ]
                        if valid_dates:
                            reader_data.append(reader)

            doc_file = generate_reregistration_statistics(reader_data, obj.month)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Просмотр статистики.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def debtors_report(request, object_id):
    if request.method == 'POST':
        obj: CreateDebtorsReport = CreateDebtorsReport.objects.get(pk=object_id)
        debtors = Debtors.objects.using('belrw-service-db').all()
        user = Reader.objects.using('belrw-user-db').get(user=request.user)

        if user.library is not None:
            debtors = debtors.filter(library=user.library)

        if obj.report_type == 'Date':
            if obj.date is not None:
                debtors = debtors.filter(refund_date__lt=obj.date)

            logger.debug(f'debtors: {debtors}')

            doc_file = generate_debt_statistics(debtors)
        else:
            if obj.reader_id is not None:
                debtors = debtors.filter(reader_id=obj.reader_id)

            doc_file = generate_debt_statistics_user(debtors)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Должники.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def debtors_report_first_class(request, object_id):
    if request.method == 'POST':
        obj: CreateDebtorsReport = CreateDebtorsReport.objects.get(pk=object_id)
        debtors = Debtors.objects.using('belrw-service-db').all()

        if obj.library is not None:
            debtors = debtors.filter(library=obj.library)

        if obj.report_type == 'Date':
            if obj.date is not None:
                debtors = debtors.filter(refund_date__lt=obj.date)


            doc_file = generate_debt_statistics(debtors)
        else:
            if obj.reader_id is not None:
                debtors = debtors.filter(reader_id=obj.reader_id)
            doc_file = generate_debt_statistics_user(debtors)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Должники.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def accountiong_statement_report(request, object_id):
    if request.method == 'POST':
        try:
            obj = AccountiongStatement.objects.get(pk=object_id)
        except AccountiongStatement.DoesNotExist:
            raise Http404("BookIncompleteEdition object does not exist")

        doc_file = generate_accountiong_statement_document(obj)

        # Отправляем файл в ответе
        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="AccountingStatementReport.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def certificate_acceptance_periodicals_report(request, object_id):
    if request.method == 'POST':
        try:
            obj = CertificateAcceptancePeriodicals.objects.get(pk=object_id)
        except CertificateAcceptancePeriodicals.DoesNotExist:
            raise Http404("BookIncompleteEdition object does not exist")

        doc_file = generate_certificate_acceptance_document(obj)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="CertificateAcceptanceReport.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def periodicals_acceptance_report(request, object_id):
    if request.method == 'POST':
        try:
            obj = PeriodicalsAcceptanceReport.objects.get(pk=object_id)
        except PeriodicalsAcceptanceReport.DoesNotExist:
            raise Http404("BookIncompleteEdition object does not exist")

        doc_file = generate_periodicals_acceptance_document(obj)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response[
            'Content-Disposition'] = f'attachment; filename="PeriodicalsAcceptanceReport_{obj.year}_{obj.month}.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def circulation_report_first_class(request, object_id):
    if request.method == 'POST':
        obj: BookCirculationReport = BookCirculationReport.objects.get(pk=object_id)
        circulations = BookCirculation.objects.using('belrw-service-db').all()

        if obj.library is not None:
            circulations = circulations.filter(library=obj.library)

        if obj.report_type == 'Date':
            if obj.year is not None:
                circulations = circulations.filter(refund_date__year__lte=obj.year)

            logger.debug(f'circulations: {circulations}')

            doc_file = generate_circulation_statistics(circulations)
        else:
            if obj.reader_id is not None:
                circulations = circulations.filter(reader_id=obj.reader_id)

            doc_file = generate_circulation_statistics_user(circulations)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Книгооборот.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def circulation_report(request, object_id):
    if request.method == 'POST':
        obj: BookCirculationReport = BookCirculationReport.objects.get(pk=object_id)
        circulations = BookCirculation.objects.using('belrw-service-db').all()
        user = Reader.objects.using('belrw-user-db').get(user=request.user)

        if user.library is not None:
            circulations = circulations.filter(library=user.library)

        if obj.report_type == 'Date':
            if obj.year is not None:
                circulations = circulations.filter(refund_date__year__lte=obj.year)

            logger.debug(f'circulations: {circulations}')

            doc_file = generate_circulation_statistics(circulations)
        else:
            if obj.reader_id is not None:
                circulations = circulations.filter(reader_id=obj.reader_id)

            doc_file = generate_circulation_statistics_user(circulations)

        response = HttpResponse(doc_file,
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Книгооборот.docx'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def typeofservice_report(request, object_id):
    if request.method == 'POST':
        obj = TypeOfService.objects.get(pk=object_id)

        doc_file = generate_typeofservice_document(obj)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="TypeOfServiceReport_{obj.year}_{obj.month}.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)


def storage_report_admin(request, object_id):
    if request.method == 'POST':
        obj = StorageReport.objects.get(pk=object_id)

        doc_file = generate_storage_report_document(obj)

        response = HttpResponse(
            doc_file,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="TypeOfServiceReport_{obj.year}_{obj.month}.docx"'
        return response
    else:
        return JsonResponse({"error": "Invalid request"}, status=400)
